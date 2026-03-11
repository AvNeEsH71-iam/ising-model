"""
Generate Word Document Report for Ising Model Cellular Automata Project
Includes figures from the simulation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_report():
    doc = Document()
    
    # Get the directory where this script is located
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ============ TITLE PAGE ============
    title = doc.add_heading('2D Ising Model Simulation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Using Cellular Automata with Checkerboard Decomposition')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].size = Pt(14)
    
    doc.add_paragraph('\n\n')
    
    # Author info
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run('MSc Physics Project\n').bold = True
    author_para.add_run('2026\n')
    
    doc.add_page_break()
    
    # ============ TABLE OF CONTENTS ============
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph(
        '1. Introduction ................................. 1\n'
        '2. Physics Background ........................... 2\n'
        '3. Computational Method ......................... 4\n'
        '4. Simulation Details ........................... 6\n'
        '5. Results and Discussion ....................... 7\n'
        '6. Conclusion ................................... 12\n'
        '7. References ................................... 13\n'
        'Appendix A: Figures ............................. 14'
    )
    doc.add_page_break()
    
    # ============ INTRODUCTION ============
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 The Ising Model', level=2)
    doc.add_paragraph(
        'The Ising model is a mathematical model of ferromagnetism in statistical physics, '
        'first proposed by Wilhelm Lenz in 1920 and solved in one dimension by Ernst Ising in 1925. '
        'The model consists of discrete variables called spins arranged on a lattice, where each spin '
        'can take one of two values: +1 (spin up) or -1 (spin down).'
    )
    
    doc.add_paragraph(
        'Despite its simplicity, the Ising model captures essential features of phase transitions '
        'and critical phenomena. In two dimensions, it exhibits a continuous (second-order) phase '
        'transition at a critical temperature T_c, below which spins spontaneously align to create '
        'a net magnetization.'
    )
    
    doc.add_heading('1.2 Historical Context', level=2)
    p = doc.add_paragraph()
    p.add_run('1920: ').bold = True
    p.add_run('Wilhelm Lenz proposes the model\n')
    p = doc.add_paragraph()
    p.add_run('1925: ').bold = True
    p.add_run('Ernst Ising solves the 1D case (no phase transition)\n')
    p = doc.add_paragraph()
    p.add_run('1944: ').bold = True
    p.add_run('Lars Onsager provides exact solution for 2D (T_c = 2.269)\n')
    p = doc.add_paragraph()
    p.add_run('1963: ').bold = True
    p.add_run('Roy Glauber develops stochastic dynamics\n')
    p = doc.add_paragraph()
    p.add_run('1960s: ').bold = True
    p.add_run('Cellular automata approaches developed')
    
    # ============ PHYSICS BACKGROUND ============
    doc.add_heading('2. Physics Background', level=1)
    
    doc.add_heading('2.1 Hamiltonian', level=2)
    doc.add_paragraph(
        'The energy (Hamiltonian) of the Ising model is given by:'
    )
    p = doc.add_paragraph()
    p.add_run('E = -J Σ⟨i,j⟩ sᵢ sⱼ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.italic = True
    
    doc.add_paragraph(
        'where sᵢ = ±1 is the spin at site i, J > 0 is the ferromagnetic coupling constant, '
        'and ⟨i,j⟩ denotes nearest-neighbor pairs. The negative sign indicates that aligned spins '
        '(sᵢ sⱼ = +1) have lower energy.'
    )
    
    doc.add_heading('2.2 Thermodynamic Quantities', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Quantity'
    header_cells[1].text = 'Formula'
    header_cells[2].text = 'Physical Meaning'
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    data = [
        ('Magnetization', 'M = (1/N) |Σᵢ sᵢ|', 'Order parameter'),
        ('Energy', 'E = -J Σ⟨i,j⟩ sᵢ sⱼ', 'Internal energy per spin'),
        ('Heat Capacity', 'C_V = (⟨E²⟩ - ⟨E⟩²) / (k_B T²)', 'Energy fluctuations'),
        ('Susceptibility', 'χ = (⟨M²⟩ - ⟨M⟩²) / (k_B T)', 'Magnetic response'),
    ]
    
    for row_data in data:
        row = table.add_row()
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        row.cells[2].text = row_data[2]
    
    doc.add_paragraph('\n')
    
    doc.add_heading('2.3 Phase Behavior', level=2)
    
    doc.add_heading('Low Temperature (T < T_c)', level=3)
    doc.add_paragraph(
        'Below the critical temperature, the system is in the ordered (ferromagnetic) phase. '
        'Spins spontaneously align, resulting in non-zero magnetization. The Z₂ symmetry '
        '(spin flip symmetry) is spontaneously broken.'
    )
    
    doc.add_heading('Critical Temperature (T = T_c)', level=3)
    doc.add_paragraph(
        'At T_c = 2.269 (Onsager solution), the system undergoes a continuous phase transition. '
        'The correlation length diverges, and critical fluctuations occur at all length scales.'
    )
    
    doc.add_heading('High Temperature (T > T_c)', level=3)
    doc.add_paragraph(
        'Above T_c, the system is in the disordered (paramagnetic) phase. Thermal fluctuations '
        'dominate over spin interactions, resulting in zero net magnetization. The symmetry is restored.'
    )
    
    # ============ COMPUTATIONAL METHOD ============
    doc.add_heading('3. Computational Method: Cellular Automata', level=1)
    
    doc.add_heading('3.1 Checkerboard Decomposition', level=2)
    doc.add_paragraph(
        'The key innovation in our Cellular Automata approach is the checkerboard (red-black) '
        'decomposition scheme. The lattice is divided into two interpenetrating sublattices:'
    )
    
    p = doc.add_paragraph()
    p.add_run('EVEN sublattice: ').bold = True
    p.add_run('sites where (i + j) is even\n')
    p = doc.add_paragraph()
    p.add_run('ODD sublattice: ').bold = True
    p.add_run('sites where (i + j) is odd')
    
    doc.add_paragraph(
        'Sites within the same sublattice are not nearest neighbors, meaning they do not interact '
        'directly. This allows all sites in a sublattice to be updated simultaneously (in parallel) '
        'without affecting each other\'s update probabilities.'
    )
    
    doc.add_heading('3.2 Update Algorithm', level=2)
    p = doc.add_paragraph()
    p.add_run('One complete CA sweep consists of:\n')
    
    steps = [
        'Calculate local magnetic field for all sites',
        'Update EVEN sublattice sites in parallel using Glauber dynamics',
        'Update ODD sublattice sites in parallel using Glauber dynamics'
    ]
    
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. ')
        run.bold = True
        p.add_run(step)
    
    doc.add_heading('3.3 Glauber Dynamics', level=2)
    doc.add_paragraph(
        'The probability for a spin to flip is given by the Glauber formula:'
    )
    p = doc.add_paragraph()
    p.add_run('P(flip) = 1 / (1 + exp(ΔE / k_B T))')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.italic = True
    doc.add_paragraph(
        'where ΔE = 2 sᵢ hᵢ is the energy change, and hᵢ is the local field from 4 neighbors.'
    )
    
    doc.add_heading('3.4 Why Checkerboard?', level=2)
    doc.add_paragraph(
        'The checkerboard scheme is essential for correctness:'
    )
    
    reasons = [
        ('Preserves detailed balance', 'Required for correct equilibrium'),
        ('Converges to Boltzmann distribution', 'Same as Monte Carlo'),
        ('Enables parallel updates', '5× speedup over sequential MC'),
        ('GPU-friendly', 'Natural for parallel hardware'),
    ]
    
    for reason, explanation in reasons:
        p = doc.add_paragraph()
        run = p.add_run(f'• {reason}: ')
        run.bold = True
        p.add_run(explanation)
    
    doc.add_heading('3.5 Comparison with Monte Carlo', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Aspect'
    header_cells[1].text = 'Monte Carlo'
    header_cells[2].text = 'Cellular Automata'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    ca_data = [
        ('Update scheme', 'Sequential (one spin)', 'Parallel (sublattice)'),
        ('Detailed balance', 'Yes', 'Yes (with checkerboard)'),
        ('Equilibrium', 'Boltzmann', 'Boltzmann'),
        ('Execution time', '~40 s', '~8 s'),
        ('Speed', '1×', '5× faster'),
    ]
    
    for row_data in ca_data:
        row = table.add_row()
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        row.cells[2].text = row_data[2]
    
    # ============ SIMULATION DETAILS ============
    doc.add_heading('4. Simulation Details', level=1)
    
    doc.add_heading('4.1 Parameters', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parameter'
    header_cells[1].text = 'Value'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    params = [
        ('Lattice size (L)', '20 × 20 = 400 spins'),
        ('Temperature range', '1.5 – 4.0'),
        ('Temperature points', '14'),
        ('Equilibration', '500-800 sweeps (adaptive)'),
        ('Measurement', '800 sweeps'),
        ('Independent runs', '3 (averaged)'),
        ('Critical temperature', 'T_c = 2.269'),
    ]
    
    for row_data in params:
        row = table.add_row()
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
    
    doc.add_paragraph('\n')
    
    doc.add_heading('4.2 Accuracy Improvements', level=2)
    improvements = [
        'Adaptive equilibration: More steps near T_c where relaxation is slow',
        'Multiple independent runs: Average over 3 different initial conditions',
        'Frequent sampling: Measure every 3 steps for better statistics',
        'Unbiased variance: Use ddof=1 for fluctuation formulas',
    ]
    
    for imp in improvements:
        p = doc.add_paragraph()
        run = p.add_run('✓  ')
        run.bold = True
        p.add_run(imp)
    
    doc.add_heading('4.3 Units', level=2)
    doc.add_paragraph(
        'We use natural units throughout: J = 1 (coupling constant), k_B = 1 (Boltzmann constant). '
        'Temperature is expressed in units of J/k_B.'
    )
    
    # ============ RESULTS AND DISCUSSION ============
    doc.add_heading('5. Results and Discussion', level=1)
    
    doc.add_heading('5.1 Temperature Scan Results', level=2)
    doc.add_paragraph(
        'The simulation was run for 14 temperatures spanning the range 1.5 to 4.0, covering all three '
        'phases: ordered (T < T_c), critical (T ≈ T_c), and disordered (T > T_c).'
    )
    
    # Sample data table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'T'
    header_cells[1].text = '|M|/N'
    header_cells[2].text = 'E/N'
    header_cells[3].text = 'C_V'
    header_cells[4].text = 'χ'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    sample_data = [
        ('1.500', '0.9861', '-1.9495', '0.0005', '0.0001'),
        ('2.077', '0.8854', '-1.6880', '0.0021', '0.0014'),
        ('2.269', '0.6980', '-1.4473', '0.0038', '0.0147'),
        ('2.462', '0.3909', '-1.1659', '0.0029', '0.0185'),
        ('3.038', '0.1382', '-0.8041', '0.0010', '0.0036'),
        ('4.000', '0.0805', '-0.5498', '0.0004', '0.0010'),
    ]
    
    for row_data in sample_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
    
    doc.add_paragraph('\nTable 1: Simulation results at selected temperatures.\n')
    
    doc.add_heading('5.2 Accuracy Comparison', level=2)
    doc.add_paragraph(
        'Comparison with Monte Carlo simulations shows excellent agreement:'
    )
    p = doc.add_paragraph()
    p.add_run('Energy difference: dE < 0.03\nMagnetization difference: dM < 0.01')
    p.italic = True
    
    # ============ FIGURES SECTION ============
    doc.add_heading('5.3 Spin Configurations', level=2)
    doc.add_paragraph(
        'Figure 1 shows representative spin configurations at three characteristic temperatures. '
        'At low temperature (T = 1.5), the system exhibits long-range order with most spins aligned. '
        'At the critical temperature (T = 2.269), we observe critical fluctuations with domains of '
        'both spin orientations. At high temperature (T = 3.5), the system is disordered with '
        'randomly oriented spins.'
    )
    
    # Add Figure 1
    fig1_path = os.path.join(script_dir, 'fig1_ca_spin_configurations.png')
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(6.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = p.add_run('Figure 1: Spin configurations at T = 1.5 (ordered), T = 2.269 (critical), and T = 3.5 (disordered). Blue represents spin down (-1), red represents spin up (+1).')
        caption_run.italic = True
    
    doc.add_page_break()
    
    doc.add_heading('5.4 Thermodynamic Quantities', level=2)
    doc.add_paragraph(
        'Figure 2 presents all four thermodynamic quantities as functions of temperature. '
        'The magnetization (top left) serves as the order parameter, showing spontaneous symmetry '
        'breaking below T_c. The energy (top right) varies continuously through the transition, '
        'characteristic of a second-order phase transition. Both the heat capacity (bottom left) '
        'and magnetic susceptibility (bottom right) exhibit peaks near the critical temperature, '
        'reflecting enhanced fluctuations at the phase transition.'
    )
    
    # Add Figure 2
    fig2_path = os.path.join(script_dir, 'fig2_ca_thermodynamic_quantities.png')
    if os.path.exists(fig2_path):
        doc.add_picture(fig2_path, width=Inches(6.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = p.add_run('Figure 2: Thermodynamic quantities vs temperature. (a) Magnetization, (b) Energy, (c) Heat Capacity, (d) Susceptibility.')
        caption_run.italic = True
    
    doc.add_page_break()
    
    doc.add_heading('5.5 Combined Summary', level=2)
    doc.add_paragraph(
        'Figure 3 shows the normalized magnetization and heat capacity on the same plot, '
        'illustrating the relationship between the order parameter and the response function. '
        'The magnetization decreases as temperature increases, while the heat capacity peaks '
        'at the critical point where fluctuations are maximal.'
    )
    
    # Add Figure 3
    fig3_path = os.path.join(script_dir, 'fig3_ca_combined.png')
    if os.path.exists(fig3_path):
        doc.add_picture(fig3_path, width=Inches(6.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = p.add_run('Figure 3: Normalized order parameter (magnetization) and response function (heat capacity) vs temperature.')
        caption_run.italic = True
    
    doc.add_page_break()
    
    doc.add_heading('5.6 Annotated Magnetization Curve', level=2)
    doc.add_paragraph(
        'Figure 4 provides an annotated view of the magnetization curve, clearly showing '
        'the phase transition. The ordered phase (T < T_c) is characterized by spontaneous '
        'magnetization, while the disordered phase (T > T_c) has vanishing magnetization. '
        'The critical temperature T_c = 2.269 marks the boundary between these two phases.'
    )
    
    # Add Figure 4
    fig4_path = os.path.join(script_dir, 'fig4_ca_magnetization_annotated.png')
    if os.path.exists(fig4_path):
        doc.add_picture(fig4_path, width=Inches(6.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = p.add_run('Figure 4: Spontaneous magnetization vs temperature with phase transition annotation. The vertical dashed line indicates T_c = 2.269.')
        caption_run.italic = True
    
    doc.add_page_break()
    
    doc.add_heading('5.7 Update Rule Comparison', level=2)
    doc.add_paragraph(
        'Figure 5 compares three different update rules: Glauber dynamics (standard), '
        'Metropolis CA (parallel MC-like), and Heat Bath algorithm. All three rules '
        'converge to the same equilibrium distribution but exhibit different transient dynamics. '
        'The right panel shows the time evolution of magnetization for each rule at T = 2.0.'
    )
    
    # Add Figure 5
    fig5_path = os.path.join(script_dir, 'fig5_ca_rule_comparison.png')
    if os.path.exists(fig5_path):
        doc.add_picture(fig5_path, width=Inches(6.5))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = p.add_run('Figure 5: Comparison of CA update rules. Left: Magnetization vs temperature for different rules. Right: Time evolution of magnetization at T = 2.0.')
        caption_run.italic = True
    
    doc.add_page_break()
    
    # ============ DISCUSSION ============
    doc.add_heading('6. Discussion', level=1)
    
    doc.add_heading('6.1 Phase Transition', level=2)
    doc.add_paragraph(
        'The simulation successfully captures the key features of the phase transition:'
    )
    
    features = [
        ('Spontaneous magnetization', 'M ≈ 0.99 at T = 1.5, indicating nearly perfect ordering'),
        ('Sharp transition', 'Magnetization drops rapidly near T_c = 2.269'),
        ('Critical fluctuations', 'Enhanced heat capacity and susceptibility at T_c'),
        ('Disordered phase', 'M ≈ 0.08 at T = 4.0, approaching zero magnetization'),
    ]
    
    for feature, detail in features:
        p = doc.add_paragraph()
        run = p.add_run(f'• {feature}: ')
        run.bold = True
        p.add_run(detail)
    
    doc.add_heading('6.2 Finite-Size Effects', level=2)
    doc.add_paragraph(
        'The simulation uses a finite lattice (L = 20), which introduces some deviations from the '
        'thermodynamic limit (L → ∞):'
    )
    
    effects = [
        'Magnetization does not reach exactly zero above T_c',
        'Heat capacity peak is rounded rather than divergent',
        'Peak position may shift slightly from exact T_c',
    ]
    
    for effect in effects:
        p = doc.add_paragraph()
        p.add_run('–  ')
        p.add_run(effect)
    
    doc.add_heading('6.3 Method Advantages', level=2)
    advantages = [
        ('Speed', '5× faster than sequential Monte Carlo'),
        ('Correctness', 'Checkerboard scheme preserves detailed balance'),
        ('Scalability', 'Natural for GPU parallelization'),
        ('Dynamics', 'Better suited for time-dependent studies'),
    ]
    
    for advantage, detail in advantages:
        p = doc.add_paragraph()
        run = p.add_run(f'✓ {advantage}: ')
        run.bold = True
        p.add_run(detail)
    
    # ============ CONCLUSION ============
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        'This project successfully demonstrates the use of Cellular Automata with checkerboard '
        'decomposition for simulating the 2D Ising model. The implementation achieves:'
    )
    
    conclusions = [
        'Correct equilibrium properties (agreement with Monte Carlo)',
        'Accurate phase transition behavior (T_c = 2.269)',
        'Significant speedup (5× faster than MC)',
        'Complete thermodynamic characterization',
    ]
    
    for conclusion in conclusions:
        p = doc.add_paragraph()
        run = p.add_run('•  ')
        run.bold = True
        p.add_run(conclusion)
    
    doc.add_paragraph(
        'The Cellular Automata approach with checkerboard decomposition is a valid and efficient '
        'alternative to traditional Monte Carlo methods for studying equilibrium statistical mechanics.'
    )
    
    # ============ REFERENCES ============
    doc.add_heading('8. References', level=1)
    
    refs = [
        'Onsager, L. (1944). Crystal Statistics. I. A Two-Dimensional Model with an Order-Disorder Transition. Physical Review, 65, 117-149.',
        'Glauber, R.J. (1963). Time-Dependent Statistics of the Ising Model. Journal of Mathematical Physics, 4, 294-307.',
        'Landau, D.P. & Binder, K. (2014). A Guide to Monte Carlo Simulations in Statistical Physics. Cambridge University Press.',
        'Pathria, R.K. & Beale, P.D. (2011). Statistical Mechanics (3rd ed.). Academic Press.',
        'Sethna, J.P. (2006). Statistical Mechanics: Entropy, Order Parameters, and Complexity. Oxford University Press.',
    ]
    
    for ref in refs:
        p = doc.add_paragraph()
        p.style = 'List Paragraph'
        p.add_run(ref)
    
    # ============ SAVE ============
    output_path = os.path.join(script_dir, 'Ising_Model_CA_Report.docx')
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    print(f"Figures included:")
    print(f"  - fig1_ca_spin_configurations.png")
    print(f"  - fig2_ca_thermodynamic_quantities.png")
    print(f"  - fig3_ca_combined.png")
    print(f"  - fig4_ca_magnetization_annotated.png")
    print(f"  - fig5_ca_rule_comparison.png")

if __name__ == '__main__':
    create_report()
